######################################################################################
#
# Author: zuoguocai@126.com
#
# Function: Get daily english TO DOCX file  Version 3.0
#
# Modified Time:  2021年4月9日
# 
# Help:  need install  requests,lxml and python-docx
#        pip3 install requests lxml python-docx -i https://pypi.douban.com/simple
#
######################################################################################
#!/usr/bin/env python



import requests

# 忽略请求https，客户端没有证书警告信息
requests.packages.urllib3.disable_warnings()


from lxml import etree


from docx import Document
from docx.shared import Inches

import time
import random

url = "https://bj.wendu.com/zixun/yingyu/6697.html"

headers = {
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/89.0.4389.114 Safari/537.36'
}



# 获取表格中 所有a标签链接

r = requests.get(url,headers=headers,verify=False,timeout=120)
html = etree.HTML(r.text)


html_urls = html.xpath("//tr//a/@href")

num_count = len(html_urls)

print("总共发现:   {}句".format(num_count))



# 获取链接下内容

for i in html_urls:
  r = requests.get(i, headers=headers, verify=False,timeout=120)
  
  result_html = etree.HTML(r.content, parser=etree.HTMLParser(encoding='utf8'))

  html_data = result_html.xpath('//div[@class="article-body"]/p//text()')
  
  # 获取标题
  head = html_data[1]
  #print(head)
  
  Message = "正在处理===>" + head  + "  "+ i + "  请稍等..."
  print(Message)
  
  # 句子和问题
  juzi = '\n'.join(html_data[2:4])
  # 选项
  xuanxiang = '\n'.join(html_data[4:10])
  # 分析
  fengxi = '\n'.join(html_data[10:-4])
  
  # 合并为一篇内容，中间以换行符分隔
  content = "\n\n\n".join((juzi,xuanxiang,fengxi))
  
  #python-docx仅可使用文档中定义的样式，因此您需要将该样式添加到模板文档中才能使用它
  file_name = 'C:\\Users\\zuoguocai\\Desktop\\pachong\\docs\\' +  '何凯文每日一句.docx'

  
  
  # 把300多句写入到一个文件里，每次打开后追加新内容，需要提前建立一个word空文件
  document = Document(file_name)
  
  # 设置字体为 考研试卷专用字体 Times New Roman 
  # 选中全部，设置为宋体，选中全部，设置为Times New Roman  
  document.styles['Normal'].font.name = u'Times New Roman'
  
  # 设置标题1，为后续 在word中查看方便(视图--导航窗格)做准备，为后续在 word中 生成目录做准备(引用--目录--自动目录)
  # 或者使用考虑使用win32com.client包对于目录进行操作
  
 
  document.add_heading(head, level=1)
  
  document.add_paragraph(content)
  document.add_page_break()
  document.save(file_name)
  
  # 限制请求
  myrandom = random.randint(3,10)
  time.sleep(myrandom)
  

  
  
  
  
 



